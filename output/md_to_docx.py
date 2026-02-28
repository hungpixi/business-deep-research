"""Convert markdown business plan to DOCX with proper formatting."""
import re
import sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

def parse_md_to_docx(md_path, docx_path):
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Remove YAML frontmatter
    content = re.sub(r'^---\n.*?---\n', '', content, flags=re.DOTALL)
    content = re.sub(r'^---\r\n.*?---\r\n', '', content, flags=re.DOTALL)
    
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Heading styles
    for i in range(1, 5):
        h_style = doc.styles[f'Heading {i}']
        h_style.font.name = 'Calibri'
        h_style.font.color.rgb = RGBColor(0x1A, 0x56, 0xDB)
        if i == 1:
            h_style.font.size = Pt(22)
        elif i == 2:
            h_style.font.size = Pt(16)
        elif i == 3:
            h_style.font.size = Pt(13)
        else:
            h_style.font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    table_buffer = []
    in_code_block = False
    code_buffer = []
    
    while i < len(lines):
        line = lines[i].rstrip('\r')
        
        # Code blocks
        if line.strip().startswith('```'):
            if in_code_block:
                code_text = '\n'.join(code_buffer)
                if code_text.strip():
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    p.paragraph_format.left_indent = Cm(1)
                code_buffer = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue
        
        # Table detection
        if '|' in line and line.strip().startswith('|'):
            table_buffer.append(line)
            i += 1
            continue
        elif table_buffer:
            _add_table(doc, table_buffer)
            table_buffer = []
        
        # Headings
        if line.startswith('# '):
            text = _clean_md(line[2:].strip())
            doc.add_heading(text, level=1)
        elif line.startswith('## '):
            text = _clean_md(line[3:].strip())
            doc.add_heading(text, level=2)
        elif line.startswith('### '):
            text = _clean_md(line[4:].strip())
            doc.add_heading(text, level=3)
        elif line.startswith('#### '):
            text = _clean_md(line[5:].strip())
            doc.add_heading(text, level=4)
        elif line.startswith('> '):
            text = _clean_md(line[2:].strip())
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(text)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = _clean_md(re.sub(r'^[\s]*[-*]\s', '', line).strip())
            p = doc.add_paragraph(style='List Bullet')
            _add_formatted_run(p, text)
        elif re.match(r'^\s*\d+\.\s', line):
            text = _clean_md(re.sub(r'^\s*\d+\.\s', '', line).strip())
            p = doc.add_paragraph(style='List Number')
            _add_formatted_run(p, text)
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
        elif line.strip():
            p = doc.add_paragraph()
            _add_formatted_run(p, _clean_md(line.strip()))
        
        i += 1
    
    if table_buffer:
        _add_table(doc, table_buffer)
    
    doc.save(docx_path)
    print(f"✅ Saved: {docx_path}")

def _clean_md(text):
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`(.*?)`', r'\1', text)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    return text

def _add_formatted_run(paragraph, text):
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)

def _add_table(doc, rows):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip('|').split('|')]
        if all(re.match(r'^[-:]+$', c) for c in cells if c):
            continue
        parsed.append(cells)
    
    if len(parsed) < 1:
        return
    
    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for r_idx, row_data in enumerate(parsed):
        for c_idx, cell_text in enumerate(row_data):
            if c_idx < num_cols:
                cell = table.cell(r_idx, c_idx)
                cell.text = _clean_md(cell_text)
                if r_idx == 0:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            run.bold = True
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(9)
                        run.font.name = 'Calibri'
    
    doc.add_paragraph()

if __name__ == '__main__':
    md_file = sys.argv[1] if len(sys.argv) > 1 else r'd:\business\startup\business-deep-research\output\edtech_ai_3d_business_plan.md'
    docx_file = sys.argv[2] if len(sys.argv) > 2 else r'd:\business\startup\business-deep-research\output\edtech_ai_3d_business_plan.docx'
    parse_md_to_docx(md_file, docx_file)
